"""
Markdown-to-DOCX converter with embedded infographic images.
Handles headings, lists, tables, code blocks, callouts, blockquotes, and inline formatting.
"""

import re
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

from style_config import (
    CALLOUT_MARKER_DEFS,
    CODE_FONT,
    FONT_FAMILY,
    HEADING_STYLE_DEFS,
    IMAGE_WIDTH_INCHES,
    MARGIN_BOTTOM_INCHES,
    MARGIN_LEFT_INCHES,
    MARGIN_RIGHT_INCHES,
    MARGIN_TOP_INCHES,
    PAGE_HEIGHT_INCHES,
    PAGE_WIDTH_INCHES,
    get_docx_colors,
)

# Resolve colors at import time (docx_builder always needs python-docx)
_COLORS = get_docx_colors()
DEEP_BLUE = _COLORS["DEEP_BLUE"]
TEAL = _COLORS["TEAL"]
ORANGE = _COLORS["ORANGE"]
DARK_GRAY = _COLORS["DARK_GRAY"]
LIGHT_GRAY = _COLORS["LIGHT_GRAY"]

# Build resolved heading styles and callout markers
HEADING_STYLES = [(size, _COLORS[color_key]) for size, color_key in HEADING_STYLE_DEFS]
CALLOUT_MARKERS = {marker: _COLORS[color_key] for marker, color_key in CALLOUT_MARKER_DEFS.items()}


class BookBuilder:
    """Build illustrated DOCX books from markdown chapters with embedded images."""

    def __init__(self, footer_text: str = ""):
        self.footer_text = footer_text

    def build(
        self,
        chapters_dir: str | Path,
        chapter_order: list[str],
        image_data: dict[int, bytes],
        output_path: str | Path,
        r2_image_dir: str | Path | None = None,
    ) -> Path:
        """Build a complete illustrated DOCX from markdown chapters.

        Args:
            chapters_dir: Directory containing .md chapter files
            chapter_order: Ordered list of chapter filenames
            image_data: Dict mapping sequential index -> image bytes (Round 1)
            output_path: Where to save the .docx
            r2_image_dir: Optional directory with r2_XXX.png files (Round 2)

        Returns:
            Path to the output file
        """
        chapters_path = Path(chapters_dir)
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        self._setup_page(doc)
        self._setup_styles(doc)

        r1_counter = 0

        for i, chapter_file in enumerate(chapter_order):
            filepath = chapters_path / chapter_file
            if not filepath.exists():
                print(f"  SKIP: {chapter_file} not found")
                continue

            print(f"  Processing: {chapter_file}")
            md_text = filepath.read_text(encoding="utf-8")

            if i > 0:
                doc.add_page_break()

            r1_counter = self._process_markdown(
                doc, md_text, image_data, r1_counter, r2_image_dir
            )

        if self.footer_text:
            self._add_footer(doc)

        doc.save(str(output))
        print(f"\nBook saved to: {output}")
        return output

    def build_from_sections(
        self,
        sections: list[dict],
        images: dict[str, bytes],
        output_path: str | Path,
    ) -> Path:
        """Build DOCX from section dicts (for document upload workflow).

        Args:
            sections: List of {title, body, level?}
            images: Dict mapping section title -> image bytes
            output_path: Where to save the .docx
        """
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        self._setup_page(doc)
        self._setup_styles(doc)

        for i, section in enumerate(sections):
            title = section["title"]
            body = section.get("body", "")
            is_top = i == 0 or section.get("level", 1) == 1

            heading_para = doc.add_heading(title, level=1 if is_top else 2)
            for run in heading_para.runs:
                run.font.size = Pt(HEADING_STYLES[0][0] if is_top else HEADING_STYLES[1][0])
                run.font.bold = True
                run.font.color.rgb = HEADING_STYLES[0][1] if is_top else HEADING_STYLES[1][1]

            if title in images and images[title]:
                self._embed_image(doc, images[title])

            if body:
                for paragraph_text in body.split("\n\n"):
                    stripped = paragraph_text.strip()
                    if stripped:
                        p = doc.add_paragraph()
                        run = p.add_run(stripped)
                        run.font.name = FONT_FAMILY
                        run.font.size = Pt(11)
                        run.font.color.rgb = DARK_GRAY

        if self.footer_text:
            self._add_footer(doc)

        doc.save(str(output))
        return output

    def _setup_page(self, doc: Document):
        section = doc.sections[0]
        section.page_width = Inches(PAGE_WIDTH_INCHES)
        section.page_height = Inches(PAGE_HEIGHT_INCHES)
        section.top_margin = Inches(MARGIN_TOP_INCHES)
        section.bottom_margin = Inches(MARGIN_BOTTOM_INCHES)
        section.left_margin = Inches(MARGIN_LEFT_INCHES)
        section.right_margin = Inches(MARGIN_RIGHT_INCHES)

    def _setup_styles(self, doc: Document):
        style = doc.styles["Normal"]
        style.font.name = FONT_FAMILY
        style.font.size = Pt(11)
        style.font.color.rgb = DARK_GRAY

        for level, (size, color) in enumerate(HEADING_STYLES, start=1):
            hs = doc.styles[f"Heading {level}"]
            hs.font.name = FONT_FAMILY
            hs.font.size = Pt(size)
            hs.font.color.rgb = color
            hs.font.bold = True

    def _embed_image(self, doc: Document, img_bytes: bytes, width: float = IMAGE_WIDTH_INCHES):
        try:
            img_stream = BytesIO(img_bytes)
            Image.open(img_stream)  # validate
            img_stream.seek(0)
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(img_stream, width=Inches(width))
        except Exception as e:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(f"[IMAGE ERROR: {e}]")
            run.font.color.rgb = ORANGE

    def _add_footer(self, doc: Document):
        footer = doc.sections[-1].footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run(self.footer_text)
        run.font.size = Pt(8)
        run.font.color.rgb = LIGHT_GRAY

    def _find_image(self, desc: str, image_data: dict[int, bytes], r1_counter: int, r2_dir) -> tuple[bytes | None, int]:
        """Find image for a marker. Check Round 2 ID first, then Round 1 sequential."""
        r2_match = re.search(r'r2_(\d{3})', desc)
        if r2_match and r2_dir:
            r2_id = r2_match.group(0)
            img_path = Path(r2_dir) / f"{r2_id}.png"
            if img_path.exists():
                return (img_path.read_bytes(), r1_counter)
            return (None, r1_counter)

        if r1_counter in image_data:
            return (image_data[r1_counter], r1_counter + 1)
        return (None, r1_counter + 1)

    def _process_markdown(self, doc, md_text, image_data, r1_counter, r2_dir):
        """Convert markdown text to DOCX elements. Returns updated r1_counter."""
        lines = md_text.split("\n")
        i = 0
        in_code_block = False
        code_buffer = []
        table_rows = []

        while i < len(lines):
            line = lines[i]

            # Code blocks
            if line.strip().startswith("```"):
                if in_code_block:
                    para = doc.add_paragraph()
                    run = para.add_run("\n".join(code_buffer))
                    run.font.name = CODE_FONT
                    run.font.size = Pt(10)
                    run.font.color.rgb = DARK_GRAY
                    code_buffer = []
                    in_code_block = False
                else:
                    in_code_block = True
                i += 1
                continue

            if in_code_block:
                code_buffer.append(line)
                i += 1
                continue

            # Tables
            if "|" in line and line.strip().startswith("|"):
                cells = [c.strip() for c in line.strip().strip("|").split("|")]
                if all(set(c) <= set("- :") for c in cells):
                    i += 1
                    continue
                table_rows.append(cells)
                if i + 1 < len(lines) and "|" in lines[i + 1] and lines[i + 1].strip().startswith("|"):
                    i += 1
                    continue
                else:
                    if table_rows:
                        num_cols = max(len(r) for r in table_rows)
                        table = doc.add_table(rows=len(table_rows), cols=num_cols)
                        table.style = "Light Grid Accent 1"
                        for ri, row in enumerate(table_rows):
                            for ci, cell_text in enumerate(row):
                                if ci < num_cols:
                                    cell = table.cell(ri, ci)
                                    cell.text = cell_text
                                    for para in cell.paragraphs:
                                        for run in para.runs:
                                            run.font.size = Pt(10)
                                            run.font.name = FONT_FAMILY
                                            if ri == 0:
                                                run.bold = True
                        doc.add_paragraph()
                        table_rows = []
                    i += 1
                    continue

            stripped = line.strip()

            if not stripped:
                i += 1
                continue

            if stripped in ("---", "***", "___"):
                doc.add_paragraph()
                i += 1
                continue

            # INFOGRAPHIC markers — embed image
            infographic_match = re.search(r"\[INFOGRAPHIC:\s*(.+?)(?:\]|$)", stripped)
            if infographic_match:
                desc = infographic_match.group(1).rstrip("]").strip()
                img_bytes, r1_counter = self._find_image(desc, image_data, r1_counter, r2_dir)

                if img_bytes:
                    self._embed_image(doc, img_bytes)
                else:
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run("[INFOGRAPHIC PLACEHOLDER]")
                    run.bold = True
                    run.font.size = Pt(12)
                    run.font.color.rgb = TEAL
                    dp = doc.add_paragraph()
                    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    dr = dp.add_run(desc[:200])
                    dr.italic = True
                    dr.font.size = Pt(9)
                    dr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

                doc.add_paragraph()
                i += 1
                continue

            # Headings
            if stripped.startswith("#"):
                match = re.match(r"^(#{1,6})\s+(.+)", stripped)
                if match:
                    level = min(len(match.group(1)), 3)
                    text = match.group(2).strip()
                    doc.add_heading(text, level=level)
                    i += 1
                    continue

            # Blockquotes
            if stripped.startswith(">"):
                text = stripped.lstrip("> ").strip()
                para = doc.add_paragraph()
                run = para.add_run(text)
                run.italic = True
                run.font.color.rgb = TEAL
                i += 1
                continue

            # Callout markers
            found_callout = False
            for marker, color in CALLOUT_MARKERS.items():
                if marker in stripped:
                    rest = stripped.replace(marker, "").strip()
                    para = doc.add_paragraph()
                    mr = para.add_run(f"{marker} ")
                    mr.bold = True
                    mr.font.color.rgb = color
                    if rest:
                        para.add_run(rest).font.color.rgb = DARK_GRAY
                    found_callout = True
                    break
            if found_callout:
                i += 1
                continue

            # Checklists
            if stripped.startswith("- [ ]") or stripped.startswith("- [x]"):
                text = stripped[5:].strip()
                symbol = "+" if stripped.startswith("- [x]") else "o"
                doc.add_paragraph(f"  {symbol}  {text}", style="List Bullet")
                i += 1
                continue

            # Bullets
            if stripped.startswith("- ") or stripped.startswith("* "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
                i += 1
                continue

            # Numbered lists
            num_match = re.match(r"^(\d+)\.\s+(.+)", stripped)
            if num_match:
                doc.add_paragraph(num_match.group(2), style="List Number")
                i += 1
                continue

            # Regular paragraphs with inline formatting
            para = doc.add_paragraph()
            parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith("*") and part.endswith("*"):
                    run = para.add_run(part[1:-1])
                    run.italic = True
                else:
                    para.add_run(part)

            i += 1

        return r1_counter
